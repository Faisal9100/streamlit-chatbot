
import os
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
import asyncio
from langchain_community.vectorstores import FAISS  # Corrected FAISS import
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI
from pypdf import PdfReader
import pandas as pd
from langchain.schema import Document
import docx  # Ensure the correct docx package is being used

# Set up Google Gemini API Key
os.environ['GOOGLE_API_KEY'] = 'AIzaSyAOAmgie2kqwwCkPUU8qdMcgLTx7X_5Brk'

# Ensure the directory for uploaded files exists
upload_directory = "uploaded_files"
os.makedirs(upload_directory, exist_ok=True)

# Function to load PDF documents
def load_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF file: {e}"

# Function to load TXT files
def load_txt(file_path):
    try:
        with open(file_path, "r") as file:
            return file.read()
    except Exception as e:
        return f"Error reading TXT file: {e}"

# Function to load DOCX files
def load_docx(file_path):
    try:
        doc = docx.Document(file_path)  # Correct usage of the Document class
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"Error reading DOCX file: {e}"

# Function to load Excel files (e.g., .xlsx, .xls)
def load_excel(file_path):
    try:
        # Ensure we are using 'openpyxl' engine for reading .xlsx files
        data = pd.read_excel(file_path, engine='openpyxl', skiprows=2)  # Specify the engine explicitly
        return data.head(10).to_string()  # Show first few rows of the first sheet
    except Exception as e:
        return f"Error reading Excel file: {e}"

# Function to load CSV files (e.g., .csv)
def load_csv(file_path):
    try:
        data = pd.read_csv(file_path)
        return data.head().to_string() if not data.empty else "CSV file is empty."
    except Exception as e:
        return f"Error reading CSV file: {e}"

# Streamlit interface for uploading files
st.title("RAG Chatbot for Course Management")

# File uploader for multiple files
uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx", "xlsx", "csv"], accept_multiple_files=True)

# Initialize extracted text as an empty list
extracted_text = []

if uploaded_files:
    # Iterate over all uploaded files and process them
    for uploaded_file in uploaded_files:
        # Create a file path in the upload directory using the uploaded file's name
        file_path = os.path.join(upload_directory, uploaded_file.name)
        
        # Save the uploaded file temporarily
        with open(file_path, "wb") as f:
            f.write(uploaded_file.read())

        # Check the file extension and MIME type, process accordingly
        if uploaded_file.type == "application/pdf":
            text = load_pdf(file_path)
            extracted_text.append(text)
        elif uploaded_file.type == "text/plain":
            text = load_txt(file_path)
            extracted_text.append(text)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = load_docx(file_path)  # Now using the updated load_docx function
            extracted_text.append(text)
        elif uploaded_file.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            # If it's an Excel file, we process it using pandas
            data = load_excel(file_path)
            extracted_text.append(data)
        elif uploaded_file.type == "text/csv":
            # If it's a CSV file, we process it using pandas
            data = load_csv(file_path)
            extracted_text.append(data)
        else:
            st.warning(f"Unsupported file type: {uploaded_file.name}. Only PDF, TXT, DOCX, Excel, and CSV are supported.")

    # Combine all extracted texts into one document
    combined_text = "\n".join(extracted_text)

    # Show a preview of the extracted text
    st.subheader("Extracted Text Preview")
    st.write(combined_text[:1000])  # Show the first 1000 characters

    if combined_text.strip():  # Check if the combined text is not empty
        # Wrap the combined extracted text into a Document object
        documents = [Document(page_content=combined_text)]

        # Now split the documents into smaller chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        docs = text_splitter.split_documents(documents)

        # Check if any documents were created
        if docs:
            # Vectorization using HuggingFace embeddings
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            vectorstore = FAISS.from_documents(docs, embeddings)

            # Set up Google Gemini model for response generation
            llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash-exp")

            # Create the RAG (Retrieval-Augmented Generation) chain
            qa = RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=vectorstore.as_retriever())

            # User query input field
            query = st.text_input("Ask a question:")

            # When the user submits a query, generate the response
            if query:
                response = qa.run(query)
                st.write(response)
        else:
            st.error("No valid chunks found after splitting the documents. Please check the input text.")
    else: 
        st.error("No valid content found in the uploaded files.")


# import os
# import streamlit as st
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_huggingface import HuggingFaceEmbeddings
# import asyncio
# from langchain_community.vectorstores import FAISS  # Corrected FAISS import
# from langchain.chains import RetrievalQA
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pypdf import PdfReader
# import pandas as pd
# from langchain.schema import Document
# import docx  # Ensure the correct docx package is being used

# # Set up Google Gemini API Key
# os.environ['GOOGLE_API_KEY'] = 'AIzaSyAOAmgie2kqwwCkPUU8qdMcgLTx7X_5Brk'

# # Ensure the directory for uploaded files exists
# upload_directory = "uploaded_files"
# os.makedirs(upload_directory, exist_ok=True)

# # Function to load PDF documents
# def load_pdf(file_path):
#     try:
#         reader = PdfReader(file_path)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text()
#         return text
#     except Exception as e:
#         return f"Error reading PDF file: {e}"

# # Function to load TXT files
# def load_txt(file_path):
#     try:
#         with open(file_path, "r") as file:
#             return file.read()
#     except Exception as e:
#         return f"Error reading TXT file: {e}"

# # Function to load DOCX files
# def load_docx(file_path):
#     try:
#         doc = docx.Document(file_path)  # Correct usage of the Document class
#         text = ""
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#         return text
#     except Exception as e:
#         return f"Error reading DOCX file: {e}"

# def load_excel(file_path):
#     try:
#         # Read the Excel file with support for multiple formats (including .xlsx)
#         data = pd.read_excel(file_path, engine='openpyxl')  # Specify the engine explicitly
#         return data.head().to_string()  # Show the first few rows of the first sheet
#     except Exception as e:
#         return f"Error reading Excel file: {e}"


# # Function to load CSV files (e.g., attendance)
# def load_csv(file_path):
#     try:
#         data = pd.read_csv(file_path)
#         return data.head().to_string() if not data.empty else "CSV file is empty."
#     except Exception as e:
#         return f"Error reading CSV file: {e}"

# # Streamlit interface for uploading files
# st.title("RAG Chatbot for Course Management")

# # File uploader for multiple files
# uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx", "xlsx", "csv"], accept_multiple_files=True)

# # Initialize extracted text as an empty list
# extracted_text = []

# if uploaded_files:
#     # Iterate over all uploaded files and process them
#     for uploaded_file in uploaded_files:
#         # Create a file path in the upload directory using the uploaded file's name
#         file_path = os.path.join(upload_directory, uploaded_file.name)
        
#         # Save the uploaded file temporarily
#         with open(file_path, "wb") as f:
#             f.write(uploaded_file.read())

#         # Check the file extension and process accordingly
#         if uploaded_file.type == "application/pdf":
#             text = load_pdf(file_path)
#             extracted_text.append(text)
#         elif uploaded_file.type == "text/plain":
#             text = load_txt(file_path)
#             extracted_text.append(text)
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             text = load_docx(file_path)  # Now using the updated load_docx function
#             extracted_text.append(text)
#         elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "text/csv":
#             # If it's an Excel or CSV file, we process it using pandas
#             if uploaded_file.type == "application/vnd.ms-excel":
#                 data = load_excel(file_path)
#             else:
#                 data = load_csv(file_path)
#             # Convert the first few rows or any relevant data to text
#             extracted_text.append(data)
#         else:
#             st.warning(f"Unsupported file type: {uploaded_file.name}. Only PDF, TXT, DOCX, Excel, and CSV are supported.")

#     # Combine all extracted texts into one document
#     combined_text = "\n".join(extracted_text)

#     # Show a preview of the extracted text
#     st.subheader("Extracted Text Preview")
#     st.write(combined_text[:1000])  # Show the first 1000 characters

#     if combined_text.strip():  # Check if the combined text is not empty
#         # Wrap the combined extracted text into a Document object
#         documents = [Document(page_content=combined_text)]

#         # Now split the documents into smaller chunks
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         docs = text_splitter.split_documents(documents)

#         # Check if any documents were created
#         if docs:
#             # Vectorization using HuggingFace embeddings
#             embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
#             vectorstore = FAISS.from_documents(docs, embeddings)

#             # Set up Google Gemini model for response generation
#             llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash-exp")

#             # Create the RAG (Retrieval-Augmented Generation) chain
#             qa = RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=vectorstore.as_retriever())

#             # User query input field
#             query = st.text_input("Ask a question:")

#             # When the user submits a query, generate the response
#             if query:
#                 response = qa.run(query)
#                 st.write(response)
#         else:
#             st.error("No valid chunks found after splitting the documents. Please check the input text.")
#     else: 
#         st.error("No valid content found in the uploaded files.")



# import os
# import streamlit as st
# from langchain_community.document_loaders import TextLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from langchain.vectorstores import FAISS
# from langchain.chains import RetrievalQA
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pypdf import PdfReader
# import pandas as pd
# from langchain.schema import Document
# from io import StringIO

# # Set up Google Gemini API Key
# os.environ['GOOGLE_API_KEY'] = 'AIzaSyAOAmgie2kqwwCkPUU8qdMcgLTx7X_5Brk'

# # Streamlit interface for uploading files
# st.title("RAG Chatbot for Course Management")

# # Ensure the directory for uploaded files exists
# upload_directory = "uploaded_files"
# os.makedirs(upload_directory, exist_ok=True)

# # Function to load and display Excel files
# def load_and_display_excel(file_path):
#     try:
#         # Attempt to read the Excel file (without additional processing)
#         data = pd.read_excel(file_path)

#         # If data is successfully loaded, show the first few rows
#         if not data.empty:
#             return data.head()  # Show only the first 5 rows
#         else:
#             return "No data found in the Excel file."
#     except Exception as e:
#         return f"Error loading Excel file: {e}"

# # Function to load PDF documents
# def load_pdf(file_path):
#     reader = PdfReader(file_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     return text

# # Function to load TXT files
# def load_txt(file_path):
#     with open(file_path, "r") as file:
#         return file.read()

# # Function to load DOCX files
# def load_docx(file_path):
#     loader = TextLoader(file_path)
#     return loader.load()

# # Function to load CSV files
# def load_csv(file_path):
#     try:
#         data = pd.read_csv(file_path)
#         return data.head().to_string() if not data.empty else "CSV file is empty."
#     except Exception as e:
#         return f"Error loading CSV file: {e}"

# # File uploader for multiple files
# uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx", "xlsx", "csv"], accept_multiple_files=True)

# # Initialize extracted text as an empty list
# extracted_text = []

# if uploaded_files:
#     # Iterate over all uploaded files and process them
#     for uploaded_file in uploaded_files:
#         # Create a file path in the upload directory using the uploaded file's name
#         file_path = os.path.join(upload_directory, uploaded_file.name)
        
#         # Save the uploaded file temporarily
#         with open(file_path, "wb") as f:
#             f.write(uploaded_file.read())

#         # Check the file extension and process accordingly
#         if uploaded_file.type == "application/pdf":
#             text = load_pdf(file_path)
#             if text.strip():
#                 extracted_text.append(text)
#             else:
#                 st.warning(f"PDF file '{uploaded_file.name}' has no text content.")
#         elif uploaded_file.type == "text/plain":
#             text = load_txt(file_path)
#             if text.strip():
#                 extracted_text.append(text)
#             else:
#                 st.warning(f"Text file '{uploaded_file.name}' has no content.")
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             text = load_docx(file_path)
#             if text.strip():
#                 extracted_text.append(text)
#             else:
#                 st.warning(f"DOCX file '{uploaded_file.name}' has no content.")
#         elif uploaded_file.type == "application/vnd.ms-excel" or uploaded_file.type == "text/csv":
#             # Process Excel or CSV file
#             if uploaded_file.type == "application/vnd.ms-excel":
#                 data = load_and_display_excel(file_path)
#                 if isinstance(data, pd.DataFrame) and not data.empty:
#                     extracted_text.append(data.to_string())
#                 else:
#                     st.warning(f"Excel file '{uploaded_file.name}' has no valid content.")
#             else:
#                 data = load_csv(file_path)
#                 if data.strip():
#                     extracted_text.append(data)
#                 else:
#                     st.warning(f"CSV file '{uploaded_file.name}' has no valid content.")

#     # Combine all extracted texts into one document
#     combined_text = "\n".join(extracted_text)

#     # Show a preview of the extracted text
#     st.subheader("Extracted Text Preview")
#     st.write(combined_text[:1000])  # Show the first 1000 characters

#     if combined_text.strip():  # Check if the combined text is not empty
#         # Wrap the combined extracted text into a Document object
#         documents = [Document(page_content=combined_text)]

#         # Now split the documents into smaller chunks
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         docs = text_splitter.split_documents(documents)

#         # Check if any documents were created
#         if docs:
#             # Vectorization using HuggingFace embeddings
#             embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
#             vectorstore = FAISS.from_documents(docs, embeddings)

#             # Set up Google Gemini model for response generation
#             llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash-exp")

#             # Create the RAG (Retrieval-Augmented Generation) chain
#             qa = RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=vectorstore.as_retriever())

#             # User query input field
#             query = st.text_input("Ask a question:")

#             # When the user submits a query, generate the response
#             if query:
#                 response = qa.run(query)
#                 st.write(response)
#         else:
#             st.error("No valid chunks found after splitting the documents. Please check the input text.")
#     else:
#         st.error("No valid content found in the uploaded files.")
